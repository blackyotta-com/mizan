import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF
import textwrap
from openai import OpenAI
import streamlit as st
import hashlib

import streamlit as st
import hashlib


USERS = {
    "Osama.Zangana": "a547448b0367490e42aa78296ca189e43ed2b8c88bd91396c1f5a71682f7b143",
    "Andulrahman.Hayek": "152c7fd3ce5cd9009d52577b7117c1e3a19bd92d8472231a48fc7bce3253fdb5"
}

if "auth_user" not in st.session_state:
    st.session_state["auth_user"] = None

def login_block():
    st.markdown("### 🔒 Secure Login")
    with st.form("login_form", clear_on_submit=True):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")
        if submitted:
            hashed = hashlib.sha256(password.encode()).hexdigest()
            if username in USERS and USERS[username] == hashed:
                st.session_state["auth_user"] = username
                st.rerun()   # <---- FIXED HERE
            else:
                st.error("Invalid username or password")

# --- LOGIN LOGIC ---
if not st.session_state.get("auth_user"):
    login_block()
    st.stop()

# --- LOGOUT LOGIC ---
if st.button("Logout", key="logout_btn"):
    st.session_state["auth_user"] = None
    st.rerun()   # <---- FIXED HERE
    st.stop()

# === SESSION STATE FOR AI SUMMARY ===
if "ai_summary" not in st.session_state:
    st.session_state["ai_summary"] = ""

# === YOUR OPENAI API KEY ===
client = OpenAI(api_key="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")

st.set_page_config(page_title="AML Investigation Report", layout="wide")
st.title("First Iraqi Bank AML Investigation System")

# Four text boxes for user input
customer_name = st.text_input("Customer's Full Name:")
job_title = st.text_input("Job Title:")
organization = st.text_input("Organization or Company:")
monthly_income = st.text_input("Monthly Income:")

uploaded_file = st.file_uploader("Upload Excel File", type=["xls", "xlsx", "xlsm"])

def parse_amount(val):
    try:
        val = str(val).replace('IQD', '').replace(',', '').replace(' ', '').replace('nan', '').strip()
        return int(val) if val else 0
    except:
        return 0

def get_summaries(df):
    df = df.copy()
    df['AMOUNT_CLEAN'] = df['AMOUNT'].apply(parse_amount)
    df['FEE_CLEAN'] = df['FEE'].apply(parse_amount) if 'FEE' in df.columns else 0
    df['TRANSACTION TYPE'] = df['TRANSACTION TYPE'].astype(str).str.upper()

    # Debit, Credit, All
    debit_df = df[df['AMOUNT_CLEAN'] < 0].copy()
    credit_df = df[df['AMOUNT_CLEAN'] > 0].copy()
    all_df = df.copy()

    # Convert all debit amounts to positive for reporting (remove -)
    debit_df['AMOUNT_CLEAN'] = debit_df['AMOUNT_CLEAN'].abs()

    def summarize(table, amount_col='AMOUNT_CLEAN'):
        total_amount = table[amount_col].sum()
        summary = (
            table
            .groupby('TRANSACTION TYPE')
            .agg({
                amount_col: 'sum',
                'FEE_CLEAN': 'sum',
                'TRANSACTION TYPE': 'count'
            })
            .rename(columns={amount_col: 'Amount', 'FEE_CLEAN': 'Fee', 'TRANSACTION TYPE': 'Count'})
            .reset_index()
        )
        summary = summary[['TRANSACTION TYPE', 'Amount', 'Fee', 'Count']]
        # Percentage column
        summary['% of Total'] = summary['Amount'] / total_amount * 100
        summary['% of Total'] = summary['% of Total'].apply(lambda x: f"{x:.2f}%" if pd.notnull(x) else "0.00%")
        return summary, total_amount

    debit_summary, debit_total = summarize(debit_df)
    credit_summary, credit_total = summarize(credit_df)
    all_summary, all_total = summarize(all_df)

    # Add grand total rows (calculate before formatting)
    def add_grand_total(summary, total_amount):
        fee_total = summary['Fee'].sum()
        count_total = summary['Count'].sum()
        total_row = pd.DataFrame([{
            'TRANSACTION TYPE': 'Grand Total',
            'Amount': total_amount,
            'Fee': fee_total,
            'Count': count_total,
            '% of Total': '100.00%'
        }])
        return pd.concat([summary, total_row], ignore_index=True)

    debit_summary = add_grand_total(debit_summary, debit_total)
    credit_summary = add_grand_total(credit_summary, credit_total)
    all_summary = add_grand_total(all_summary, all_total)

    # Format numbers with commas (after all calculations)
    for s in [debit_summary, credit_summary, all_summary]:
        s['Amount'] = s['Amount'].apply(lambda x: f"{int(x):,}" if pd.notnull(x) and x != '' else "0")
        s['Fee'] = s['Fee'].apply(lambda x: f"{int(x):,}" if pd.notnull(x) and x != '' else "0")
        s['Count'] = s['Count'].apply(lambda x: f"{int(x):,}" if pd.notnull(x) and x != '' else "0")

    return debit_summary, credit_summary, all_summary

def style_header_and_last_row(df, last_color="#27AE60"):
    header_style = 'background-color: #222; color: #fff; font-weight: bold;'
    last_row_style = f'background-color: {last_color}; color: #fff; font-weight: bold;'
    default_style = ''

    style_df = pd.DataFrame(default_style, index=df.index, columns=df.columns)
    style_df.iloc[0, :] = header_style
    style_df.iloc[-1, :] = last_row_style

    return df.style.apply(lambda _: style_df, axis=None)

def set_cell_background(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def generate_docx(
    debit, credit, all_summary, comments,
    customer_name, job_title, organization, monthly_income,
    ai_conclusion=None
):
    doc = Document()
    doc.add_heading('AML Investigation Report', 0)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_heading("Customer Information", level=1)
    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    fields = [
        ("Customer's Full Name", customer_name),
        ("Job Title", job_title),
        ("Organization/Company", organization),
        ("Monthly Income", monthly_income)
    ]
    for idx, (field, value) in enumerate(fields):
        info.cell(idx, 0).text = field
        info.cell(idx, 1).text = value

    def add_table(doc, title, df):
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = col
        for idx, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                val = row[col]
                cells[i].text = str(val)
        doc.add_paragraph()  # Blank line

    add_table(doc, "Debit Transactions Summary", debit)
    add_table(doc, "Credit Transactions Summary", credit)
    add_table(doc, "All Transactions Summary", all_summary)

    doc.add_heading('Analyst Comments', level=1)
    doc.add_paragraph(comments)

    if ai_conclusion:
        doc.add_heading('AI Analysis Conclusion', level=1)
        doc.add_paragraph(ai_conclusion)

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def clean_multiline_text(text, wrap=90, max_word=45):
    out = []
    for line in text.splitlines():
        for bad in ["\u200b", "\u202a", "\t"]:
            line = line.replace(bad, " ")
        words = line.split(" ")
        safe_words = []
        for word in words:
            while len(word) > max_word:
                safe_words.append(word[:max_word])
                word = word[max_word:]
            if word:
                safe_words.append(word)
        safe_line = " ".join(safe_words)
        out.extend(textwrap.wrap(safe_line, wrap) or [" "])
    return out

def export_pdf(
    debit, credit, all_summary, comments,
    customer_name, job_title, organization, monthly_income
):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "AML Investigation Report", ln=1, align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 8, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=1)
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Customer Information", ln=1)
    pdf.set_font("Arial", "", 11)
    fields = [
        ("Customer's Full Name", customer_name),
        ("Job Title", job_title),
        ("Organization/Company", organization),
        ("Monthly Income", monthly_income)
    ]
    for field, value in fields:
        pdf.cell(50, 7, f"{field}:", border=0)
        pdf.cell(0, 7, str(value), ln=1, border=0)
    pdf.ln(2)

    def add_table(title, df, last_color=(0,0,0)):
        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, title, ln=1)
        pdf.set_font("Arial", "B", 10)
        col_widths = [40, 28, 28, 22, 30]
        headers = list(df.columns)
        pdf.set_fill_color(34, 34, 34)
        pdf.set_text_color(255, 255, 255)
        for i, col in enumerate(headers):
            pdf.cell(col_widths[i], 8, str(col), border=1, align="C", fill=True)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", "", 10)
        for idx, row in df.iterrows():
            fill = False
            if idx == len(df)-1:
                pdf.set_fill_color(*last_color)
                pdf.set_text_color(255,255,255)
                fill = True
            for i, col in enumerate(headers):
                cell_val = str(row[col])
                pdf.cell(col_widths[i], 7, cell_val, border=1, align="C", fill=fill)
            pdf.ln()
            pdf.set_fill_color(255,255,255)
            pdf.set_text_color(0,0,0)
        pdf.ln(3)

    add_table("Debit Transactions Summary", debit, last_color=(231,76,60))
    add_table("Credit Transactions Summary", credit, last_color=(39,174,96))
    add_table("All Transactions Summary", all_summary, last_color=(46,134,193))

    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Analyst Comments", ln=1)
    pdf.set_font("Arial", "", 11)
    for line in clean_multiline_text(comments, wrap=90):
        pdf.multi_cell(0, 7, line)
    pdf.ln(2)

    pdf_bytes = BytesIO(pdf.output(dest='S'))
    pdf_bytes.seek(0)
    return pdf_bytes

def create_ai_prompt(customer_name, job_title, organization, monthly_income, debit_summary, credit_summary, all_summary):
    def table_text(df, title):
        lines = [f"{title}:"]
        for _, row in df.iterrows():
            if row['TRANSACTION TYPE'] != 'Grand Total':
                lines.append(
                    f"- {row['TRANSACTION TYPE']}: Amount {row['Amount']} ({row['% of Total']})"
                )
        return "\n".join(lines)
    
    prompt = f"""
Customer Information:
Full Name: {customer_name}
Job Title: {job_title}
Organization/Company: {organization}
Monthly Income: {monthly_income}

{table_text(debit_summary, "Debit Transactions Summary")}

{table_text(credit_summary, "Credit Transactions Summary")}

{table_text(all_summary, "All Transactions Summary")}

Write a clear, structured AML/transaction monitoring case summary with:
- Preliminary Assessment
- Key statistics (as in the above tables)
- Conclusion Depending on below Points, Mention which Point is fits the Case, Mention the Point and Tell it's Fits it: 

1-Funds transferred to and from the same account on the same day or within a relatively short period of time.

2- Purchases that exceed the predetermined POS transaction limit.

3- A customer conducting multiple POS purchases within a single day.

4- A customer conducting multiple POS purchases within 30 days.

5- Flagging as suspicious any top-up transactions ranging between 3,925,000 IQD and 4,000,000 IQD, or between $2,975 and $3,025.

6- Flagging as suspicious any top-up transactions for the same card totaling 1,250,000 IQD or more, received from different agents within short time intervals.

This is because individuals subject to fraud often have funds transferred to the same card by different agents. The fraudster publishes the card details, resulting in a "many-to-one" transfer scenario, but via top-ups and agent deposits.

7- Flagging as suspicious any deposit or transfer transaction exceeding 1,000,000 IQD, followed by cash withdrawal or fund transfer operations from the recipient card.

In such cases, the fraudster typically withdraws or spends the money immediately from ATMs or retail channels before the victim can discover the fraud and report it to the company or bank to block the fraudster's card.

8- The customer's monthly income exceeds the expected income declared in the KYC documentation.(clear AML risk reasoning and suspicious features if present)

Format: Use clear headings, bullet points, and professional, only answer the Conclusion in Arabic Language.
"""
    return prompt

if uploaded_file:
    xls = pd.ExcelFile(uploaded_file)
    first_sheet = xls.sheet_names[0]
    df = pd.read_excel(uploaded_file, sheet_name=first_sheet)
    if not set(['AMOUNT', 'TRANSACTION TYPE']).issubset(df.columns):
        st.error("The uploaded file does not contain required columns. Check your sheet format.")
    else:
        debit_summary, credit_summary, all_summary = get_summaries(df)
        st.markdown("### Investigation Summaries")

        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("#### <span style='color:#FFE135;background:#222;border-radius:4px;padding:4px;'>Debit Transactions Summary</span>", unsafe_allow_html=True)
            st.table(style_header_and_last_row(debit_summary, "#E74C3C"))  # Red

        with col2:
            st.markdown("#### <span style='color:#44FF44;background:#222;border-radius:4px;padding:4px;'>Credit Transactions Summary</span>", unsafe_allow_html=True)
            st.table(style_header_and_last_row(credit_summary, "#27AE60"))  # Green

        with col3:
            st.markdown("#### <span style='color:#2E86C1;background:#222;border-radius:4px;padding:4px;'>All Transactions Summary</span>", unsafe_allow_html=True)
            st.table(style_header_and_last_row(all_summary, "#2E86C1"))  # Blue

        # === TOP 5 TRANSACTIONS FOR EACH TYPE ===
        st.markdown("## Top 5 Transactions for Each Transaction Type")
        if "TRANSACTION TYPE" in df.columns and "AMOUNT" in df.columns:
            # Build a DATE_TIME column if possible
            if "DATE" in df.columns and "TIME" in df.columns:
                df["DATE_TIME"] = df["DATE"].astype(str) + " " + df["TIME"].astype(str)
            elif "DATE" in df.columns:
                df["DATE_TIME"] = df["DATE"].astype(str)
            else:
                df["DATE_TIME"] = ""
            for ttype in df["TRANSACTION TYPE"].unique():
                st.markdown(f"**{ttype}**")
                ttype_df = df[df["TRANSACTION TYPE"] == ttype].copy()
                ttype_df["AMOUNT_ABS"] = ttype_df["AMOUNT"].apply(lambda x: abs(parse_amount(x)))
                top5 = ttype_df.sort_values("AMOUNT_ABS", ascending=False).head(5)
                # Always try to show Date+Time, Amount, Name, Beneficiary, Counterparty, Description
                display_cols = [col for col in ["DATE_TIME", "AMOUNT", "NAME", "BENEFICIARY", "COUNTERPARTY", "DESCRIPTION"] if col in top5.columns]
                if display_cols:
                    st.dataframe(top5[display_cols])
                else:
                    st.dataframe(top5)
        else:
            st.warning("Cannot display top transactions: missing 'TRANSACTION TYPE' or 'AMOUNT' columns.")

        # === TOP 10 NAMES BY FREQUENCY (with exclusions) ===
        name_col = None
        for possible in ["NAME", "BENEFICIARY", "COUNTERPARTY"]:
            if possible in df.columns:
                name_col = possible
                break

        st.markdown("## Top 10 Names in Transactions")
        if name_col:
            # Names/phrases to exclude (partial match, case-insensitive)
            exclude_names = [
                "Commission for P2P",
                "Service Store Purchase",
                "Card Selling Company",
                "Fastpay Erbil Branch",
                "Bank Operation"
            ]
            def not_excluded(val):
                val_lower = str(val).lower()
                return not any(ex.lower() in val_lower for ex in exclude_names)
            filtered_names = df[df[name_col].apply(not_excluded)]
            name_counts = filtered_names[name_col].value_counts().head(10).reset_index()
            name_counts.columns = [name_col, "Count"]
            st.dataframe(name_counts)
        else:
            st.info("No name column ('NAME', 'BENEFICIARY', or 'COUNTERPARTY') found to display top 10 names.")

        comments = st.text_area("Analyst Comments/Findings", "")

        # Analyze with AI button - avoid double output
        analyze_pressed = st.button("Analyze The Case")

        if analyze_pressed:
            with st.spinner("Contacting AI Analyst..."):
                prompt = create_ai_prompt(
                    customer_name, job_title, organization, monthly_income,
                    debit_summary, credit_summary, all_summary
                )
                try:
                    response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional AML compliance analyst. Respond with professional English, clear summary, and conclusion."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=800,
                        temperature=0.2,
                    )
                    st.session_state["ai_summary"] = response.choices[0].message.content
                except Exception as e:
                    st.session_state["ai_summary"] = f"Error: {e}"

        if analyze_pressed or st.session_state["ai_summary"]:
            st.markdown("### AI Analysis Summary")
            st.info(st.session_state["ai_summary"])

        # Word export (without AI)
        if st.button("Export Report as Word (.docx)"):
            docx_file = generate_docx(
                debit_summary, credit_summary, all_summary, comments,
                customer_name, job_title, organization, monthly_income
            )
            st.success("Report generated! Click below to download:")
            st.download_button(
                label="Download Word Report",
                data=docx_file,
                file_name=f"AML_Investigation_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            )
        # PDF export (without AI summary)
        if st.button("Export Report as PDF"):
            pdf_file = export_pdf(
                debit_summary, credit_summary, all_summary, comments,
                customer_name, job_title, organization, monthly_income
            )
            st.success("PDF generated! Click below to download:")
            st.download_button(
                label="Download PDF Report",
                data=pdf_file,
                file_name=f"AML_Investigation_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                mime="application/pdf"
            )

        # FULL WORD export with AI conclusion
        if st.button("Download Full Word Report with Conclusion"):
            if not st.session_state["ai_summary"]:
                with st.spinner("Contacting AI Analyst..."):
                    prompt = create_ai_prompt(
                        customer_name, job_title, organization, monthly_income,
                        debit_summary, credit_summary, all_summary
                    )
                    try:
                        response = client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "system", "content": "You are a professional AML compliance analyst. Respond with professional English, clear summary, and conclusion."},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=800,
                            temperature=0.2,
                        )
                        st.session_state["ai_summary"] = response.choices[0].message.content
                    except Exception as e:
                        st.session_state["ai_summary"] = f"Error: {e}"

            docx_file = generate_docx(
                debit_summary, credit_summary, all_summary, comments,
                customer_name, job_title, organization, monthly_income,
                ai_conclusion=st.session_state["ai_summary"]
            )
            st.success("Full Word report generated with conclusion! Click below to download:")
            st.download_button(
                label="Download Full Word Report",
                data=docx_file,
                file_name=f"AML_Investigation_Report_{datetime.now().strftime('%Y%m%d_%H%M')}_FULL.docx"
            )

else:
    st.info("Upload an Excel file to begin.")

st.markdown("---\n*No data is stored. Close the tab to clear everything. --- Credit : Blackyotta Company")


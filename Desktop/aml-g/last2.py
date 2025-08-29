# ============================
# First Iraqi Bank AML System
# Extended with Manager/Analyst Dashboards
# ============================

import os, sqlite3, time
from contextlib import closing
from passlib.hash import bcrypt
import matplotlib.pyplot as plt
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import calendar
import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import datetime
from fpdf import FPDF
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# ---------- Page config (call once) ----------
st.set_page_config(page_title="AML Investigation Report", layout="wide")

# Optional Arabic shaping
try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    _ARABIC_OK = True
except Exception:
    _ARABIC_OK = False

# OpenAI client (optional)
try:
    client = OpenAI(api_key="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")   # requires OPENAI_API_KEY in env
except Exception:
    client = None

# ----------------------------
# DB CONFIG
# ----------------------------
DB_PATH = "aml.db"

def db():
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def column_exists(conn, table, col):
    cur = conn.cursor()
    cur.execute(f"PRAGMA table_info({table})")
    return any(row[1] == col for row in cur.fetchall())

def init_db():
    with closing(db()) as conn:
        cur = conn.cursor()
        cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            full_name TEXT NOT NULL,
            job_title TEXT,
            role TEXT CHECK(role IN ('manager','analyst')) NOT NULL DEFAULT 'analyst',
            employee_id TEXT,
            manager_id TEXT,
            password_hash TEXT NOT NULL,
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            last_login_at TEXT,
            last_seen_at REAL
        );
        """)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS cases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            analyst_id INTEGER NOT NULL,
            case_title TEXT NOT NULL DEFAULT '',
            status TEXT CHECK(status IN ('open','solved')) NOT NULL DEFAULT 'open',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            closed_at TEXT,
            FOREIGN KEY(analyst_id) REFERENCES users(id)
        );
        """)
        conn.commit()

def migrate_db_if_needed():
    with closing(db()) as conn:
        cur = conn.cursor()

        # Ensure users table exists (and columns)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            full_name TEXT NOT NULL,
            job_title TEXT,
            role TEXT CHECK(role IN ('manager','analyst')) NOT NULL DEFAULT 'analyst',
            employee_id TEXT,
            manager_id TEXT,
            password_hash TEXT NOT NULL,
            is_active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            last_login_at TEXT,
            last_seen_at REAL
        );
        """)
        conn.commit()

        for col_def in [
            ("full_name", "TEXT NOT NULL DEFAULT ''"),
            ("job_title", "TEXT"),
            ("employee_id", "TEXT"),
            ("manager_id", "TEXT"),
            ("role", "TEXT DEFAULT 'analyst'"),
            ("last_seen_at", "REAL"),
            ("is_active", "INTEGER NOT NULL DEFAULT 1"),
        ]:
            cur.execute("PRAGMA table_info(users)")
            if not any(row[1] == col_def[0] for row in cur.fetchall()):
                cur.execute(f"ALTER TABLE users ADD COLUMN {col_def[0]} {col_def[1]};")
        conn.commit()

        # Ensure cases table exists
        cur.execute("""
        CREATE TABLE IF NOT EXISTS cases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            analyst_id INTEGER NOT NULL,
            case_title TEXT NOT NULL DEFAULT '',
            status TEXT CHECK(status IN ('open','solved')) NOT NULL DEFAULT 'open',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            closed_at TEXT,
            FOREIGN KEY(analyst_id) REFERENCES users(id)
        );
        """)
        conn.commit()

        # Add 'case_title' if missing
        cur.execute("PRAGMA table_info(cases)")
        cols = [r[1] for r in cur.fetchall()]
        if "case_title" not in cols:
            cur.execute("ALTER TABLE cases ADD COLUMN case_title TEXT NOT NULL DEFAULT ''")
            conn.commit()


def seed_manager_if_empty():
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM users")
        if c.fetchone()[0] == 0:
            c.execute("""INSERT INTO users (username, full_name, role, manager_id, password_hash)
                         VALUES (?,?,?,?,?)""",
                      ("manager", "Default Manager", "manager", "MGR-0001", bcrypt.hash("ChangeMe!123")))
            conn.commit()

# ----------------------------
# AUTH
# ----------------------------
if "auth_user" not in st.session_state:
    st.session_state["auth_user"] = None
if "ai_summary" not in st.session_state:
    st.session_state["ai_summary"] = ""

def auth_login(username, password):
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("SELECT id, username, role, password_hash, is_active, full_name, job_title, employee_id, manager_id FROM users WHERE username=?", (username,))
        row = c.fetchone()
        if not row or row[4] == 0:
            return None
        if bcrypt.verify(password, row[3]):
            c.execute("UPDATE users SET last_login_at=CURRENT_TIMESTAMP WHERE id=?", (row[0],))
            conn.commit()
            return {
                "id": row[0], "username": row[1], "role": row[2],
                "full_name": row[5], "job_title": row[6],
                "employee_id": row[7], "manager_id": row[8]
            }
    return None

def heartbeat(user_id):
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("UPDATE users SET last_seen_at=? WHERE id=?", (time.time(), user_id))
        conn.commit()

def is_online(last_seen_at, threshold=90):
    if not last_seen_at:
        return False
    return (time.time() - last_seen_at) <= threshold

def login_block():
    st.markdown("### 🔒 Secure Login")
    with st.form("login_form", clear_on_submit=True):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")
        if submitted:
            user = auth_login(username, password)
            if user:
                st.session_state["auth_user"] = user
                st.rerun()
            else:
                st.error("Invalid credentials or inactive account.")

# ----------------------------
# CASE HELPERS
# ----------------------------
def open_case_for_current(analyst_id, case_title):
    title_clean = (case_title or "").strip()
    with closing(db()) as conn:
        c = conn.cursor()
        try:
            c.execute(
                "INSERT INTO cases (analyst_id, case_title, status) VALUES (?, ?, 'open')",
                (analyst_id, title_clean)
            )
        except sqlite3.OperationalError as e:
            # If the column still doesn't exist, add it and retry once
            if "no column named case_title" in str(e):
                c.execute("ALTER TABLE cases ADD COLUMN case_title TEXT NOT NULL DEFAULT ''")
                conn.commit()
                c.execute(
                    "INSERT INTO cases (analyst_id, case_title, status) VALUES (?, ?, 'open')",
                    (analyst_id, title_clean)
                )
            else:
                raise
        conn.commit()


def mark_case_solved(case_id):
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("UPDATE cases SET status='solved', closed_at=CURRENT_TIMESTAMP WHERE id=?", (case_id,))
        conn.commit()

def my_stats(analyst_id):
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("SELECT COUNT(*) FROM cases WHERE analyst_id=? AND status='solved'", (analyst_id,))
        solved = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM cases WHERE analyst_id=? AND status='open'", (analyst_id,))
        open_ = c.fetchone()[0]
        return solved, open_

def my_open_cases(analyst_id):
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("""
            SELECT id, case_title, created_at
            FROM cases
            WHERE analyst_id=? AND status='open'
            ORDER BY created_at ASC
        """, (analyst_id,))
        return c.fetchall()

def team_leaderboard():
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("""
            SELECT u.id, u.username, u.full_name, u.job_title, u.role, u.employee_id, u.manager_id,
                   u.is_active, u.last_seen_at,
                   SUM(CASE WHEN cases.status='solved' THEN 1 ELSE 0 END) AS solved_count
            FROM users u
            LEFT JOIN cases ON cases.analyst_id = u.id
            GROUP BY u.id
            ORDER BY solved_count DESC, u.username ASC
        """)
        return c.fetchall()
def get_analyst_monthly_cases(analyst_id):
    """Get monthly solved cases for a specific analyst"""
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("""
            SELECT 
                strftime('%Y-%m', closed_at) as month,
                COUNT(*) as cases_solved
            FROM cases 
            WHERE analyst_id = ? AND status = 'solved' AND closed_at IS NOT NULL
            GROUP BY strftime('%Y-%m', closed_at)
            ORDER BY month DESC
            LIMIT 12
        """, (analyst_id,))
        return c.fetchall()

def get_all_analysts_monthly_cases():
    """Get monthly solved cases for all analysts (Manager view)"""
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("""
            SELECT 
                u.username,
                u.full_name,
                strftime('%Y-%m', c.closed_at) as month,
                COUNT(*) as cases_solved
            FROM cases c
            JOIN users u ON c.analyst_id = u.id
            WHERE c.status = 'solved' AND c.closed_at IS NOT NULL
            GROUP BY u.id, u.username, u.full_name, strftime('%Y-%m', c.closed_at)
            ORDER BY month DESC, cases_solved DESC
        """)
        return c.fetchall()

def create_analyst_monthly_chart(analyst_id):
    """Create bar chart for analyst's monthly performance"""
    data = get_analyst_monthly_cases(analyst_id)
    
    if not data:
        st.info("No solved cases data available for chart.")
        return
    
    # Prepare data for plotting
    months = [row[0] for row in data]
    cases = [row[1] for row in data]
    
    # Convert month strings to readable format
    month_labels = []
    for month_str in months:
        year, month = month_str.split('-')
        month_name = calendar.month_abbr[int(month)]
        month_labels.append(f"{month_name} {year}")
    
    # Create Plotly bar chart
    fig = go.Figure(data=[
        go.Bar(
            x=month_labels[::-1],  # Reverse to show oldest first
            y=cases[::-1],
            marker_color='#27AE60',
            text=cases[::-1],
            textposition='auto',
        )
    ])
    
    fig.update_layout(
        title="📊 My Monthly Case Closures",
        xaxis_title="Month",
        yaxis_title="Cases Solved",
        showlegend=False,
        height=400,
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
    )
    
    st.plotly_chart(fig, use_container_width=True)

def create_manager_team_chart():
    """Create bar chart for manager showing all analysts' performance"""
    data = get_all_analysts_monthly_cases()
    
    if not data:
        st.info("No team performance data available for chart.")
        return
    
    # Process data for plotting
    df_data = []
    for row in data:
        username, full_name, month, cases = row
        year, month_num = month.split('-')
        month_name = calendar.month_abbr[int(month_num)]
        month_label = f"{month_name} {year}"
        
        df_data.append({
            'Analyst': f"{full_name} ({username})",
            'Month': month_label,
            'Cases': cases,
            'Sort_Month': month  # For sorting
        })
    
    if not df_data:
        st.info("No data to display.")
        return
    
    df = pd.DataFrame(df_data)
    
    # Get last 6 months for better visualization
    recent_months = df['Sort_Month'].unique()
    recent_months = sorted(recent_months, reverse=True)[:6]
    df_filtered = df[df['Sort_Month'].isin(recent_months)]
    
    # Create grouped bar chart
    fig = px.bar(
        df_filtered, 
        x='Month', 
        y='Cases', 
        color='Analyst',
        title="📈 Team Performance - Monthly Case Closures",
        text='Cases'
    )
    
    fig.update_traces(texttemplate='%{text}', textposition='outside')
    fig.update_layout(
        height=500,
        xaxis_title="Month",
        yaxis_title="Cases Solved",
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        showlegend=True
    )
    
    st.plotly_chart(fig, use_container_width=True)

def get_analyst_performance_summary(analyst_id):
    """Get performance summary for analyst"""
    with closing(db()) as conn:
        c = conn.cursor()
        
        # Get current month stats
        c.execute("""
            SELECT COUNT(*) FROM cases 
            WHERE analyst_id = ? AND status = 'solved' 
            AND strftime('%Y-%m', closed_at) = strftime('%Y-%m', 'now')
        """, (analyst_id,))
        current_month = c.fetchone()[0]
        
        # Get last month stats
        c.execute("""
            SELECT COUNT(*) FROM cases 
            WHERE analyst_id = ? AND status = 'solved' 
            AND strftime('%Y-%m', closed_at) = strftime('%Y-%m', 'now', '-1 month')
        """, (analyst_id,))
        last_month = c.fetchone()[0]
        
        # Get average monthly performance
        c.execute("""
            SELECT AVG(monthly_count) FROM (
                SELECT COUNT(*) as monthly_count
                FROM cases 
                WHERE analyst_id = ? AND status = 'solved' AND closed_at IS NOT NULL
                GROUP BY strftime('%Y-%m', closed_at)
            )
        """, (analyst_id,))
        avg_monthly = c.fetchone()[0] or 0
        
        return current_month, last_month, round(avg_monthly, 1)
# ----------------------------
# DATA / EDD HELPERS
# ----------------------------
def parse_amount(val):
    try:
        val = str(val).replace('IQD', '').replace(',', '').replace(' ', '').replace('nan', '').strip()
        return int(float(val)) if val else 0
    except:
        return 0

def get_summaries(df):
    df = df.copy()
    df['AMOUNT_CLEAN'] = df['AMOUNT'].apply(parse_amount)
    df['FEE_CLEAN'] = df['FEE'].apply(parse_amount) if 'FEE' in df.columns else 0
    df['TRANSACTION TYPE'] = df['TRANSACTION TYPE'].astype(str).str.upper()

    debit_df = df[df['AMOUNT_CLEAN'] < 0].copy()
    credit_df = df[df['AMOUNT_CLEAN'] > 0].copy()
    all_df = df.copy()
    debit_df['AMOUNT_CLEAN'] = debit_df['AMOUNT_CLEAN'].abs()

    def summarize(table, amount_col='AMOUNT_CLEAN'):
        total_amount = table[amount_col].sum()
        summary = (
            table
            .groupby('TRANSACTION TYPE')
            .agg({
                amount_col: 'sum',
                'FEE_CLEAN': 'sum',
                'TRANSACTION TYPE': 'count'
            })
            .rename(columns={amount_col: 'Amount', 'FEE_CLEAN': 'Fee', 'TRANSACTION TYPE': 'Count'})
            .reset_index()
        )
        summary = summary[['TRANSACTION TYPE', 'Amount', 'Fee', 'Count']]
        summary['% of Total'] = summary['Amount'] / total_amount * 100 if total_amount != 0 else 0
        summary['% of Total'] = summary['% of Total'].apply(lambda x: f"{x:.2f}%" if pd.notnull(x) else "0.00%")
        return summary, total_amount

    debit_summary, debit_total = summarize(debit_df)
    credit_summary, credit_total = summarize(credit_df)
    all_summary, all_total = summarize(all_df)

    def add_grand_total(summary, total_amount):
        fee_total = summary['Fee'].sum()
        count_total = summary['Count'].sum()
        total_row = pd.DataFrame([{
            'TRANSACTION TYPE': 'Grand Total',
            'Amount': total_amount,
            'Fee': fee_total,
            'Count': count_total,
            '% of Total': '100.00%' if total_amount != 0 else '0.00%'
        }])
        return pd.concat([summary, total_row], ignore_index=True)

    debit_summary = add_grand_total(debit_summary, debit_total)
    credit_summary = add_grand_total(credit_summary, credit_total)
    all_summary = add_grand_total(all_summary, all_total)

    for s in [debit_summary, credit_summary, all_summary]:
        s['Amount'] = s['Amount'].apply(lambda x: f"{int(x):,}" if pd.notnull(x) and x != '' else "0")
        s['Fee'] = s['Fee'].apply(lambda x: f"{int(x):,}" if pd.notnull(x) and x != '' else "0")
        s['Count'] = s['Count'].apply(lambda x: f"{int(x):,}" if pd.notnull(x) and x != '' else "0")

    return debit_summary, credit_summary, all_summary

def style_header_and_last_row(df, last_color="#27AE60"):
    header_style = 'background-color: #222; color: #fff; font-weight: bold;'
    last_row_style = f'background-color: {last_color}; color: #fff; font-weight: bold;'
    default_style = ''
    style_df = pd.DataFrame(default_style, index=df.index, columns=df.columns)
    style_df.iloc[0, :] = header_style
    style_df.iloc[-1, :] = last_row_style
    return df.style.apply(lambda _: style_df, axis=None)

def create_ai_prompt(customer_name, job_title, organization, monthly_income, annual_salary, annual_income, mobile, country, city, nationality, cus_num, economic_sector, country_of_residence, resident_flag, politically_exposed_flag, debit_summary, credit_summary, all_summary):
    def table_text(df, title):
        lines = [f"{title}:"]
        for _, row in df.iterrows():
            if row['TRANSACTION TYPE'] != 'Grand Total':
                lines.append(f"- {row['TRANSACTION TYPE']}: Amount {row['Amount']} ({row['% of Total']})")
        return "\n".join(lines)
    prompt = f"""
Customer Information:
Full Name: {customer_name}
Job Title: {job_title}
Organization/Company: {organization}
Monthly Income: {monthly_income}
Annual Salary: {annual_salary}
Annual Income: {annual_income}
Mobile: {mobile}
Country: {country}
City: {city}
Nationality: {nationality}
Customer Number: {cus_num}
Economic Sector: {economic_sector}
Country of Residence: {country_of_residence}
Resident Flag: {resident_flag}
Politically Exposed Flag: {politically_exposed_flag}

{table_text(debit_summary, "Debit Transactions Summary")}

{table_text(credit_summary, "Credit Transactions Summary")}

{table_text(all_summary, "All Transactions Summary")}

Write a clear, structured AML/transaction monitoring case summary with:
- Preliminary Assessment
- Key statistics (as in the above tables)
- Conclusion Depending on below Flags, Mention which Flag is fits the Case, Mention the Flag and Tell it's Fits it: if the Case Aligned with 3 or More Flags Please mention in the End " EDD Report is Recommended for this Case " if less then 3 Flags Aligned with the Case Please mentuion " False Positive or EDD is not Recommended for the Case "

1- The customer provides information that is obviously incomplete or false, such as faking their place of residence.
2- The value of the transaction is not in line with the customer's usual activities or does not match the information available from other sources.
3- Lack of clarity regarding the repeated nature of transactions (deposits, withdrawals, transfers, etc.) or insufficient explanations for these transactions.
4- The customer acts as an intermediary for other people or collects money from multiple individuals into a single account.
5- The customer’s activity appears inconsistent with their declared income or economic situation.
6- The customer maintains a balance in an electronic wallet that does not match their declared income, with repeated top-ups and withdrawals that do not correspond to their known business activities.
7- Repeated fund transfers from multiple sources to a single card, with funds quickly withdrawn or spent.
8- voidance of disclosure about place of residence and the nature of business or economic activity.
9- The use of the same phone number or address by multiple clients, or frequent changes in address or phone number with the purpose of misleading.
10- Multiple accounts linked to one phone number or address without a clear reason.
11- Establishment of companies with forged or fake documents in order to conduct transactions under the name of those companies for purposes of money laundering or terrorist financing.
12- Clients providing financial institutions with fake or suspicious documentation or multiple addresses and phone numbers.
14- Cash deposits by persons who are not clearly related to the account holder.
15- Deposits by several parties or the receipt of checks or transfers from multiple parties.
16- Unusual deposits or deposits from sources with no clear relationship to the account holder.
17- The use of ATMs or bank branches by the customer to avoid interacting directly with bank staff.
18- Receiving foreign transfers from countries suspected of having terrorist elements or being supportive of such organizations.
19- Multiple transfers received and then re-sent to other parties, or through several intermediaries, or transferred on behalf of individuals.
20- The customer’s account being used to receive and send funds on behalf of several people without a clear relationship.
21- Repeated reception or sending of remittances to/from the same or several people without clear reason.
22- if customer receive repeated amount with round value such as 2,000,000 or 1,500,000 or 500,000
23- Flagging as suspicious any top-up transactions ranging between 3,925,000 IQD and 4,000,000 IQD, or between $2,975 and $3,025.
24- Flagging as suspicious any top-up transactions for the same card totaling 1,250,000 IQD or more, received from different agents within short time intervals.
25- Flagging as suspicious any deposit or transfer transaction exceeding 1,000,000 IQD, followed by cash withdrawal or fund transfer operations from the recipient card.

Format: Use clear headings, bullet points, and professional,if the Case Aligned with 3 or More Flags Please mention in the End " EDD Report is Recommended for this Case " if less then 3 Flags Aligned with the Case Please mention " False Positive Alert "

"""
    return prompt

def ensure_datetime_cols(df):
    if "DATE" in df.columns and "TIME" in df.columns:
        return df.assign(DATE_TIME=df["DATE"].astype(str) + " " + df["TIME"].astype(str))
    elif "DATE" in df.columns:
        return df.assign(DATE_TIME=df["DATE"].astype(str))
    else:
        return df.assign(DATE_TIME="")

def col_exists(df, names):
    for n in names:
        if n in df.columns:
            return n
    return None

def top10_names_table(df):
    name_col = col_exists(df, ["NAME", "BENEFICIARY", "COUNTERPARTY"])
    if not name_col:
        return pd.DataFrame(columns=["الاسم", "عدد العمليات"])
    exclude_names = [
        "Commission for P2P", "Service Store Purchase", "Card Selling Company",
        "Fastpay Erbil Branch", "Bank Operation"
    ]
    def not_excluded(val):
        val_lower = str(val).lower()
        return not any(ex.lower() in val_lower for ex in exclude_names)
    filtered = df[df[name_col].apply(not_excluded)]
    vc = filtered[name_col].value_counts().head(10).reset_index()
    vc.columns = ["الاسم", "عدد العمليات"]
    return vc

def get_top5_credit_transactions(df):
    d = df.copy()
    d["AMOUNT_VAL"] = d["AMOUNT"].apply(parse_amount)
    d = d[d["AMOUNT_VAL"] > 0]
    d = ensure_datetime_cols(d)
    cols = [c for c in ["DATE_TIME", "AMOUNT", "TRANSACTION TYPE", "NAME", "BENEFICIARY", "COUNTERPARTY", "DESCRIPTION"] if c in d.columns]
    d["ABS"] = d["AMOUNT_VAL"].abs()
    return d.sort_values("ABS", ascending=False).head(5)[cols]

def get_top5_credit_p2p(df):
    d = df.copy()
    d["AMOUNT_VAL"] = d["AMOUNT"].apply(parse_amount)
    d = d[(d["AMOUNT_VAL"] > 0) & (d["TRANSACTION TYPE"].str.upper().str.contains("P2P"))]
    d = ensure_datetime_cols(d)
    d["ABS"] = d["AMOUNT_VAL"].abs()
    cols = [c for c in ["DATE_TIME", "AMOUNT", "NAME", "BENEFICIARY", "COUNTERPARTY", "DESCRIPTION"] if c in d.columns]
    return d.sort_values("ABS", ascending=False).head(5)[cols]

def get_cash_deposits(df):
    d = df.copy()
    d["AMOUNT_VAL"] = d["AMOUNT"].apply(parse_amount)
    d = d[(d["AMOUNT_VAL"] > 0) &
          (d["TRANSACTION TYPE"].str.upper().str.contains("CASH")) &
          (d["TRANSACTION TYPE"].str.upper().str.contains("DEPOSIT"))]
    d = ensure_datetime_cols(d)
    d = d.rename(columns={
        "DATE_TIME": "التاريخ",
        "AMOUNT": "المبلغ /دينار",
        "DESCRIPTION": "الغاية",
        "BRANCH": "الفرع الذي تم الإيداع من خلاله",
        "SOURCE_OF_FUNDS": "مصدر الاموال حسب اشعار الايداع النقدي والمستندات المعززة ان وجدت",
        "RELATIONSHIP": "علاقة المودع مع صاحب الحساب",
        "NAME": "اسم المودع بالكامل"
    })
    cols = ["التاريخ","المبلغ /دينار","الغاية","الفرع الذي تم الإيداع من خلاله",
            "مصدر الاموال حسب اشعار الايداع النقدي والمستندات المعززة ان وجدت",
            "علاقة المودع مع صاحب الحساب","اسم المودع بالكامل"]
    for c in cols:
        if c not in d.columns:
            d[c] = ""
    d["ABS"] = d["AMOUNT_VAL"].abs()
    d = d.sort_values("ABS", ascending=False)
    return d[cols].head(15)

def get_cash_withdrawals(df):
    d = df.copy()
    d["AMOUNT_VAL"] = d["AMOUNT"].apply(parse_amount)
    d = d[d["AMOUNT_VAL"] < 0]
    d = ensure_datetime_cols(d)
    d = d[d["TRANSACTION TYPE"].str.upper().str.contains("CASH") & d["TRANSACTION TYPE"].str.upper().str.contains("WITHDRAW")]
    d["AMOUNT_POS"] = d["AMOUNT_VAL"].abs()
    d = d.rename(columns={"DATE_TIME": "التاريخ", "AMOUNT_POS": "المبلغ/ دينار", "DESCRIPTION": "الغاية"})
    cols = ["التاريخ", "المبلغ/ دينار", "الغاية"]
    for c in cols:
        if c not in d.columns:
            d[c] = ""
    return d[cols].sort_values("المبلغ/ دينار", ascending=False).head(15)

def get_top5_debit_p2p(df):
    d = df.copy()
    d["AMOUNT_VAL"] = d["AMOUNT"].apply(parse_amount)
    d = d[(d["AMOUNT_VAL"] < 0) & (d["TRANSACTION TYPE"].str.upper().str.contains("P2P"))]
    d = ensure_datetime_cols(d)
    d["ABS"] = d["AMOUNT_VAL"].abs()
    cols = [c for c in ["DATE_TIME", "AMOUNT", "NAME", "BENEFICIARY", "COUNTERPARTY", "DESCRIPTION"] if c in d.columns]
    return d.sort_values("ABS", ascending=False).head(5)[cols]

def get_debit_p2p_all(df):
    d = df.copy()
    d["AMOUNT_VAL"] = d["AMOUNT"].apply(parse_amount)
    d = d[(d["AMOUNT_VAL"] < 0) & (d["TRANSACTION TYPE"].str.upper().str.contains("P2P"))]
    d = ensure_datetime_cols(d)
    cols = [c for c in ["DATE_TIME", "AMOUNT", "NAME", "BENEFICIARY", "COUNTERPARTY", "DESCRIPTION"] if c in d.columns]
    return d[cols].head(30)

# ----------------------------
# DOCX helpers (RTL/Arabic)
# ----------------------------
PDF_WORD_FONT = "DejaVu Sans"

def _rtl_cell(cell, font_name=PDF_WORD_FONT, font_size=11):
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    try:
        p._element.get_or_add_pPr().set(qn('w:bidi'), '1')
    except Exception:
        pass
    for run in p.runs:
        run.font.name = font_name
        try:
            run._element.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
        run.font.size = Pt(font_size)

def add_df_as_table_docx(doc, df, headers_ar=None, col_widths=None, fallback_rows=3, font_name=PDF_WORD_FONT, font_size=11):
    if df is not None and not df.empty:
        cols = list(df.columns)
        display_headers = headers_ar if (headers_ar and len(headers_ar) == len(cols)) else cols
        table = doc.add_table(rows=len(df) + 1, cols=len(cols)); table.style = "Table Grid"
        for j, h in enumerate(display_headers):
            cell = table.cell(0, j); cell.text = str(h); _rtl_cell(cell, font_name, font_size)
        for i in range(len(df)):
            for j, c in enumerate(cols):
                val = "" if pd.isna(df.iloc[i][c]) else str(df.iloc[i][c])
                cell = table.cell(i + 1, j); cell.text = val; _rtl_cell(cell, font_name, font_size)
        if col_widths:
            for j, w in enumerate(col_widths):
                for i in range(len(df) + 1):
                    table.cell(i, j).width = Inches(w)
        return

    cols = headers_ar if headers_ar and len(headers_ar) > 0 else [""]
    table = doc.add_table(rows=fallback_rows + 1, cols=len(cols)); table.style = "Table Grid"
    for j, h in enumerate(cols):
        cell = table.cell(0, j); cell.text = str(h); _rtl_cell(cell, font_name, font_size)
    for i in range(1, fallback_rows + 1):
        for j in range(len(cols)):
            cell = table.cell(i, j); cell.text = "............."; _rtl_cell(cell, font_name, font_size)
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(fallback_rows + 1):
                table.cell(i, j).width = Inches(w)

def para_ar(doc, text, bold=False, font_name=PDF_WORD_FONT, font_size=12):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r = p.add_run(text); r.bold = bold; r.font.name = font_name
    try:
        r._element.rPr.rFonts.set(qn('w:cs'), font_name)
    except Exception:
        pass
    r.font.size = Pt(font_size)
    try:
        p._element.get_or_add_pPr().set(qn('w:bidi'), '1')
    except Exception:
        pass
    return p

# ----------------------------
# EDD DOCX builder
# ----------------------------
def generate_edd_docx(customer_name="", job_title="", organization="", monthly_income="",
                      annual_salary="", annual_income="", mobile="", country="", city="",
                      nationality="", cus_num="", economic_sector="", country_of_residence="",
                      resident_flag="", politically_exposed_flag="",
                      individual_or_corporate="", df=None):
    PDF_WORD_FONT = "DejaVu Sans"
    def cell_rtl(cell):
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        for run in p.runs:
            run.font.name = PDF_WORD_FONT
            try:
                run._element.rPr.rFonts.set(qn('w:cs'), PDF_WORD_FONT)
            except Exception:
                pass
            run.font.size = Pt(11)
        try:
            p._element.get_or_add_pPr().set(qn('w:bidi'), '1')
        except Exception:
            pass
    def para(text, bold=False):
        p = Document().add_paragraph()  # dummy, replaced below; keep signature
        return p

    doc = Document()
    style = doc.styles['Normal']; style.font.name = PDF_WORD_FONT; style.font.size = Pt(12)
    try: style._element.rPr.rFonts.set(qn('w:cs'), PDF_WORD_FONT)
    except Exception: pass

    def para_ar2(text, bold=False):
        p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = p.add_run(text); run.font.name = PDF_WORD_FONT
        try: run._element.rPr.rFonts.set(qn('w:cs'), PDF_WORD_FONT)
        except Exception: pass
        run.bold = bold; run.font.size = Pt(12)
        try: p._element.get_or_add_pPr().set(qn('w:bidi'), '1')
        except Exception: pass
        return p

    para_ar2("نموذج التحليل المالي", bold=True)
    para_ar2("اولا:بيانات التعرف على العميل:(customer profile)")
    para_ar2(f'· يحتفظ "{customer_name or "............."}" ({nationality or "............."} الجنسية) بحساب/ حسابات ............... لدى المصرف ......../ فرع .......  منذ ../.../........، وقد أظهرت بيانات التعرف على العميل للمذكور ما يلي:-')
    para_ar2("· تبعاً لنموذج طلب فتح الحساب بتاريخ .../.../........ :")

    para_ar2("المعلومات الشخصية:")
    personal_info_table = doc.add_table(rows=9, cols=2); personal_info_table.style = "Table Grid"
    personal_labels = ["الاسم :", "الجنسية :", "الجنسيات الأخرى :", "مكان الإقامة :",
                       "الوظيفة التي يشغلها/طبيعة العمل/الجهة :", "مصادر دخل أخرى :",
                       "الدخل الشهري/ السنوي :", "المصارف التي يتعامل معها :", "معلومات أخرى :"]
    personal_values = [customer_name or ".............",
                       nationality or ".............",
                       ".............", ".............",
                       f"{job_title or '.............'} / {organization or '.............'}",
                       ".............",
                       f"{monthly_income or '.............'} / {annual_salary or annual_income or '.............'}",
                       ".............", "............."]
    for i, (label, value) in enumerate(zip(personal_labels, personal_values)):
        c0 = personal_info_table.cell(i, 0); c0.text = label; cell_rtl(c0)
        c1 = personal_info_table.cell(i, 1); c1.text = value; cell_rtl(c1)
    doc.add_paragraph("")

    para_ar2("· وثائق إثبات الشخصية:")
    docs_table = doc.add_table(rows=5, cols=5); docs_table.style = "Table Grid"
    docs_headers = ["نوع الوثيقة","رقم","تاريخ الإصدار","تاريخ الانتهاء","ملاحظات"]
    for j, h in enumerate(docs_headers):
        c = docs_table.cell(0, j); c.text = h; cell_rtl(c)
    for row in [
        ["هوية موحدة",".............",".............",".............","............."],
        ["جواز الســـــــفر",".............",".............",".............","............."],
        ["وثيقة الإقامة",".............",".............",".............","............."],
        ["وثائق اخرى",".............",".............",".............","............."],
    ]:
        r = docs_table.add_row().cells
        for j, val in enumerate(row):
            r[j].text = val; cell_rtl(r[j])
    doc.add_paragraph("")

    para_ar2("· بيانات الشخص/ الاشخاص المفوضين بالتوقيع على الحساب (ان وجد) مع بيان صلاحيات التفويض:-"); para_ar2(".............")
    para_ar2("· بيانات الشخص/ الاشخاص المفوضين بالتوقيع على الحساب (ان وجد) مع تحديد صلاحيات التفويض:-"); para_ar2(".............")
    para_ar2("· بالرجوع إلى قوائم الحظر الدولية و المحلية (...................EU, UK, OFAC, UN,) تبين الاتي:-"); para_ar2(".............")
    para_ar2("· بالرجوع إلى محركات البحث على الشبكة العنكبوتية مثل (Google) تبين ان:-"); para_ar2(".............")

    para_ar2("ثانياً: التحليل المالي للفترة منذ تاريخ فتح الحساب ولغاية ../.../........ (او آخر سنتين او الفترة  المطلوبة):-")
    para_ar2("ملاحظة: في حال وجود أي تعامل مالي بعملة غير الدينار العراقي يجب إدراج ما يُعادلها بالدينار أينما وردت في التحليل المالي.")
    doc.add_paragraph("")

    para_ar2("· بيانات الحسابات الدائنة:")
    cred_accounts = doc.add_table(rows=4, cols=8); cred_accounts.style = "Table Grid"
    acc_headers = ["الرقم","رقم الحساب","نوع الحساب","العملة","الفرع","تاريخ فتح الحساب","الرصيد بتاريخ ......","ملاحظات"]
    for j, h in enumerate(acc_headers):
        c = cred_accounts.cell(0, j); c.text = h; cell_rtl(c)
    for i in range(1,4):
        for j in range(8):
            c = cred_accounts.cell(i, j); c.text = "............."; cell_rtl(c)
    doc.add_paragraph("")

    para_ar2("· وجود/ عدم وجود بطاقات صراف أو ائتمان مع العميل، ..."); para_ar2(".............")
    para_ar2("· وجود/ عدم وجود بطاقات صراف أو ائتمان فرعية، ..."); para_ar2(".............")
    para_ar2("· وجود/ عدم وجود صناديق حديدية ..."); para_ar2(".............")
    doc.add_paragraph("")

    para_ar2("· لدى دراسة حساب/ حسابات المدعو ...")
    para_ar2("العمليات المالية الدائنة:")
    para_ar2("بلغ مجموع العمليات المالية الدائنة (...........) ... تفاصيلها كما يلي:-")

    para_ar2("· إيداعات نقدية :")
    cash_dep_df = get_cash_deposits(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, cash_dep_df,
        headers_ar=["التاريخ","المبلغ /دينار","الغاية","الفرع الذي تم الإيداع من خلاله",
                    "مصدر الاموال حسب اشعار الايداع النقدي والمستندات المعززة ان وجدت",
                    "علاقة المودع مع صاحب الحساب","اسم المودع بالكامل"],
        fallback_rows=6, font_name=PDF_WORD_FONT)

    doc.add_paragraph("")
    para_ar2("· كما يوضح الجدول أدناه الأشخاص الأكثر إيداعاً نقدياً في الحساب:-")
    top10_df = top10_names_table(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, top10_df,
        headers_ar=["اسم المودع بالكامل","عدد الايداعات"], fallback_rows=2, font_name=PDF_WORD_FONT)

    doc.add_paragraph("")
    para_ar2("· والجدول أدناه يوضح مرسلي أكبر عدد / مبالغ حوالات واردة:")
    top5_credit_df = get_top5_credit_transactions(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, top5_credit_df,
        headers_ar=["التاريخ/الوقت","المبلغ","نوع العملية","الاسم","المستفيد","الطرف المقابل","الوصف"],
        fallback_rows=2, font_name=PDF_WORD_FONT)

    doc.add_paragraph("")
    para_ar2("· تحويلات داخلية واردة:")
    top5_credit_p2p_df = get_top5_credit_p2p(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, top5_credit_p2p_df,
        headers_ar=["التاريخ/الوقت","المبلغ","الاسم","المستفيد","الطرف المقابل","الوصف"],
        fallback_rows=2, font_name=PDF_WORD_FONT)

    para_ar2("العمليات المالية المدينة:")
    para_ar2("بلغ مجموع العمليات المالية المدينة (...........) ... تفاصيلها كما يلي:-")

    para_ar2("· السحوبات النقدية")
    cash_wd_df = get_cash_withdrawals(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, cash_wd_df,
        headers_ar=["التاريخ","المبلغ/ دينار","الغاية"], fallback_rows=3, font_name=PDF_WORD_FONT)

    para_ar2("· سحب محلي صراف آلي:"); para_ar2(".............")
    para_ar2("· سحب دولي صراف آلي:"); para_ar2(".............")
    atm_int = doc.add_table(rows=4, cols=3); atm_int.style = "Table Grid"
    for j, h in enumerate(["التاريخ","المبلغ مقوم بالدينار","الدولة التي تم بها السحب"]):
        c = atm_int.cell(0, j); c.text = h; cell_rtl(c)
    for i in range(1,4):
        for j in range(3):
            c = atm_int.cell(i, j); c.text = "............."; cell_rtl(c)

    doc.add_paragraph("")
    para_ar2("· مشتريات من خلال بطاقات الصراف/ بطاقات الائتمان:")
    purchases = doc.add_table(rows=4, cols=4); purchases.style = "Table Grid"
    for j, h in enumerate(["التاريخ","المبلغ بالعملة الاصلية","المبلغ مقوم بالدينار","اسم الجهة / الدولة"]):
        c = purchases.cell(0, j); c.text = h; cell_rtl(c)
    for i in range(1,4):
        for j in range(4):
            c = purchases.cell(i, j); c.text = "............."; cell_rtl(c)

    doc.add_paragraph("")
    para_ar2("· شيكات مسحوبة من الحساب نقداً:")
    cashcheque = doc.add_table(rows=4, cols=4); cashcheque.style = "Table Grid"
    for j, h in enumerate(["التاريخ","المبلغ/ دينار","اسم المستفيد والعلاقة","الغاية"]):
        c = cashcheque.cell(0, j); c.text = h; cell_rtl(c)
    for i in range(1,4):
        for j in range(4):
            c = cashcheque.cell(i, j); c.text = "............."; cell_rtl(c)

    doc.add_paragraph("")
    para_ar2("· شيكات مسحوبة من الحساب / مقاصة:")
    chequeclr = doc.add_table(rows=4, cols=5); chequeclr.style = "Table Grid"
    for j, h in enumerate(["الغاية","التاريخ","المبلغ/ دينار","اسم المستفيد","بنك المستفيد"]):
        c = chequeclr.cell(0, j); c.text = h; cell_rtl(c)
    for i in range(1,4):
        for j in range(5):
            c = chequeclr.cell(i, j); c.text = "............."; cell_rtl(c)

    doc.add_paragraph("")
    para_ar2("· والجدول أدناه يوضح المستفيد صاحب أكبر عدد/ مبالغ حوالات صادرة من الحساب:-")
    top5_debit_p2p_df = get_top5_debit_p2p(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, top5_debit_p2p_df,
        headers_ar=["التاريخ/الوقت","المبلغ","الاسم","المستفيد","الطرف المقابل","الوصف"],
        fallback_rows=2, font_name=PDF_WORD_FONT)

    para_ar2("· تحويلات داخلية صادرة:")
    debit_p2p_all_df = get_debit_p2p_all(df) if isinstance(df, pd.DataFrame) else pd.DataFrame()
    add_df_as_table_docx(doc, debit_p2p_all_df,
        headers_ar=["التاريخ/الوقت","المبلغ","الاسم","المستفيد","الطرف المقابل","الوصف"],
        fallback_rows=2, font_name=PDF_WORD_FONT)

    para_ar2("أي عمليات مالية أو تفاصيل اخرى ذات اهمية:-"); para_ar2("............."); para_ar2(".............")
    para_ar2("تقييم البنك لمدى تناسب حركة حسابات العميل مع طبيعة عمله / نشاطه:"); para_ar2(".............")

    buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

# ----------------------------
# PDF (Arabic-capable) + EDD PDF builder
# ----------------------------
class PDF_AR(FPDF):
    def __init__(self, *args, **kwargs):
        super().__init__(orientation='P', unit='mm', format='A4')
        try:
            self.add_font('DejaVu', '', fname='DejaVuSans.ttf', uni=True)
            self.set_font('DejaVu', '', 12)
            self._has_dejavu = True
        except:
            self.set_font('Helvetica', '', 12)
            self._has_dejavu = False

    def ar(self, text):
        if not isinstance(text, str):
            text = str(text)
        if _ARABIC_OK:
            try:
                reshaped = arabic_reshaper.reshape(text)
                return get_display(reshaped)
            except Exception:
                return text
        return text

    def hline(self, h=6):
        self.ln(h)

    def write_rtl(self, txt, h=6):
        w = self.w - self.l_margin - self.r_margin
        self.multi_cell(w, h, self.ar(txt), align='R')

    def table(self, headers, rows, col_widths=None, line_h=7):
        if not headers:
            return
        ncol = len(headers)
        if not col_widths:
            col_widths = [ (self.w - self.l_margin - self.r_margin) / ncol ] * ncol
        self.set_fill_color(230,230,230)
        for j,h in enumerate(headers):
            self.cell(col_widths[j], line_h, self.ar(str(h)), border=1, ln=0, align='C', fill=True)
        self.ln(line_h)
        for r in rows:
            for j in range(ncol):
                celltxt = "" if j>=len(r) else ("" if pd.isna(r[j]) else str(r[j]))
                self.cell(col_widths[j], line_h, self.ar(celltxt), border=1, ln=0, align='C')
            self.ln(line_h)

def df_to_rows(df):
    if df is None or df.empty:
        return []
    return [[("" if pd.isna(v) else str(v)) for v in row] for _, row in df.iterrows()]

def export_edd_pdf(customer_name="", job_title="", organization="", monthly_income="",
                   annual_salary="", annual_income="", mobile="", country="", city="",
                   nationality="", cus_num="", economic_sector="", country_of_residence="",
                   resident_flag="", politically_exposed_flag="",
                   individual_or_corporate="", df=None):
    pdf = PDF_AR(); pdf.add_page()
    pdf.set_font_size(14); pdf.write_rtl("نموذج التحليل المالي"); pdf.hline(3)
    pdf.set_font_size(12)
    pdf.write_rtl("اولا:بيانات التعرف على العميل:(customer profile)")
    pdf.write_rtl(f'· يحتفظ "{customer_name or "............."}" ({nationality or "............."} الجنسية) بحساب/ حسابات ............... لدى المصرف ......../ فرع .......  منذ ../.../........، وقد أظهرت بيانات التعرف على العميل للمذكور ما يلي:-')
    pdf.write_rtl("· تبعاً لنموذج طلب فتح الحساب بتاريخ .../.../........ :"); pdf.hline(2)

    headers = ["القيمة","الحقل"]
    rows = [
        [customer_name or ".............", "الاسم :"],
        [nationality or ".............", "الجنسية :"],
        [".............", "الجنسيات الأخرى :"],
        [".............", "مكان الإقامة :"],
        [f"{job_title or '.............'} / {organization or '.............'}", "الوظيفة التي يشغلها/طبيعة العمل/الجهة :"],
        [".............", "مصادر دخل أخرى :"],
        [f"{monthly_income or '.............'} / {annual_salary or annual_income or '.............'}", "الدخل الشهري/ السنوي :"],
        [".............", "المصارف التي يتعامل معها :"],
        [".............", "معلومات أخرى :"],
    ]
    pdf.table(headers, rows, col_widths=[90,90]); pdf.hline(3)

    pdf.write_rtl("· وثائق إثبات الشخصية:")
    pdf.table(["ملاحظات","تاريخ الانتهاء","تاريخ الإصدار","رقم","نوع الوثيقة"], [
        [".............",".............",".............",".............","هوية موحدة"],
        [".............",".............",".............",".............","جواز الســـــــفر"],
        [".............",".............",".............",".............","وثيقة الإقامة"],
        [".............",".............",".............",".............","وثائق اخرى"],
    ], col_widths=[36,36,36,36,36]); pdf.hline(2)

    pdf.write_rtl("· بيانات الشخص/ الاشخاص المفوضين ..."); pdf.write_rtl(".............")
    pdf.write_rtl("· قوائم الحظر الدولية والمحلية ..."); pdf.write_rtl(".............")
    pdf.write_rtl("· نتائج محركات البحث ..."); pdf.write_rtl("............."); pdf.hline(2)

    pdf.write_rtl("ثانياً: التحليل المالي ..."); pdf.write_rtl("ملاحظة: ..."); pdf.hline(2)

    pdf.write_rtl("· بيانات الحسابات الدائنة:")
    pdf.table(["ملاحظات","الرصيد بتاريخ ......","تاريخ فتح الحساب","الفرع","العملة","نوع الحساب","رقم الحساب","الرقم"],
              [["","","","","","","","1"],["","","","","","","","2"],["","","","","","","","3"]]); pdf.hline(2)

    pdf.write_rtl("· لدى دراسة الحسابات ..."); pdf.write_rtl("العمليات المالية الدائنة:")
    pdf.write_rtl("بلغ مجموع العمليات المالية الدائنة (...........) ..."); pdf.hline(2)

    pdf.write_rtl("· إيداعات نقدية :")
    cash_dep_df = get_cash_deposits(df) if df is not None else pd.DataFrame()
    pdf.table(["التاريخ","المبلغ /دينار","الغاية","الفرع","مصدر الاموال","العلاقة","اسم المودع"], df_to_rows(cash_dep_df),
              col_widths=[25,25,35,25,30,25,25]); pdf.hline(2)

    pdf.write_rtl("· الأشخاص الأكثر إيداعاً:")
    top10_df = top10_names_table(df) if df is not None else pd.DataFrame()
    pdf.table(["اسم المودع بالكامل","عدد الايداعات"], df_to_rows(top10_df)); pdf.hline(2)

    pdf.write_rtl("· مرسلو أكبر حوالات واردة:")
    top5_credit_df = get_top5_credit_transactions(df) if df is not None else pd.DataFrame()
    pdf.table(["التاريخ/الوقت","المبلغ","نوع العملية","الاسم","المستفيد","الطرف المقابل","الوصف"], df_to_rows(top5_credit_df)); pdf.hline(2)

    pdf.write_rtl("· تحويلات داخلية واردة:")
    top5_credit_p2p_df = get_top5_credit_p2p(df) if df is not None else pd.DataFrame()
    pdf.table(["التاريخ/الوقت","المبلغ","الاسم","المستفيد","الطرف المقابل","الوصف"], df_to_rows(top5_credit_p2p_df)); pdf.hline(2)

    pdf.write_rtl("العمليات المالية المدينة:"); pdf.write_rtl("بلغ مجموع العمليات المالية المدينة ..."); pdf.hline(2)

    pdf.write_rtl("· السحوبات النقدية")
    cash_wd_df = get_cash_withdrawals(df) if df is not None else pd.DataFrame()
    pdf.table(["التاريخ","المبلغ/ دينار","الغاية"], df_to_rows(cash_wd_df), col_widths=[40,40,100]); pdf.hline(2)

    pdf.write_rtl("· سحب محلي صراف آلي:"); pdf.write_rtl(".............")
    pdf.write_rtl("· سحب دولي صراف آلي:"); pdf.write_rtl(".............")
    pdf.write_rtl("· تفاصيل السحوبات الدولية:"); pdf.table(["التاريخ","المبلغ مقوم بالدينار","الدولة"],
              [[".............",".............","............."]]*3); pdf.hline(2)

    pdf.write_rtl("· مشتريات بالبطاقات:"); pdf.table(["التاريخ","المبلغ بالعملة الاصلية","المبلغ مقوم بالدينار","اسم الجهة / الدولة"],
              [["............."]*4]*3); pdf.hline(2)

    pdf.write_rtl("· شيكات نقداً:"); pdf.table(["التاريخ","المبلغ/ دينار","اسم المستفيد والعلاقة","الغاية"],
              [["............."]*4]*3); pdf.hline(2)

    pdf.write_rtl("· شيكات / مقاصة:"); pdf.table(["الغاية","التاريخ","المبلغ/ دينار","اسم المستفيد","بنك المستفيد"],
              [["............."]*5]*3); pdf.hline(2)

    pdf.write_rtl("· أكبر مستفيد حوالات صادرة:")
    top5_debit_p2p_df = get_top5_debit_p2p(df) if df is not None else pd.DataFrame()
    pdf.table(["التاريخ/الوقت","المبلغ","الاسم","المستفيد","الطرف المقابل","الوصف"], df_to_rows(top5_debit_p2p_df)); pdf.hline(2)

    pdf.write_rtl("· تحويلات داخلية صادرة:")
    debit_p2p_all_df = get_debit_p2p_all(df) if df is not None else pd.DataFrame()
    pdf.table(["التاريخ/الوقت","المبلغ","الاسم","المستفيد","الطرف المقابل","الوصف"], df_to_rows(debit_p2p_all_df)); pdf.hline(2)

    pdf.write_rtl("أي عمليات أخرى ذات أهمية:"); pdf.write_rtl(".............")
    pdf.write_rtl("تقييم البنك لملاءمة حركة الحساب:"); pdf.write_rtl(".............")

    out = BytesIO(); out.write(pdf.output(dest="S").encode("latin1")); out.seek(0); return out

# ----------------------------
# OPERATIONS PAGE (ALL UI INSIDE)
# ----------------------------
def operations_page():
    st.header("AML / EDD Operations")
    st.title("First Iraqi Bank AML Investigation System")

    # Case Controls - Only for Analysts
    current_user = st.session_state["auth_user"]
    if current_user["role"] == "analyst":
        st.subheader("Case Controls")
        with st.form("create_named_case_form"):
            default_name = f"Case {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            case_title = st.text_input("New Case Title", value=default_name)
            create_case = st.form_submit_button("Open New Case")
        if create_case:
            title_clean = (case_title or "").strip()
            if not title_clean:
                st.error("Please enter a case title.")
            else:
                me = st.session_state["auth_user"]
                open_case_for_current(me["id"], title_clean)
                st.success(f"Opened new case: {title_clean}")

    # Inputs
    # Customer Type Selection
    customer_type = st.selectbox("Select Customer Type",
                                ["Individual/Personal", "Business/Corporate"],
                                key="customer_type_selector")

    customer_name = st.text_input("Customer's Full Name:")
    job_title = st.text_input("Job Title:")
    organization = st.text_input("Organization or Company:")
    monthly_income = st.text_input("Monthly Income:")
    annual_salary = st.text_input("Annual Salary:")

    # Show Annual Income only for Business/Corporate customers
    if customer_type == "Business/Corporate":
        annual_income = st.text_input("Annual Income:")
    else:
        annual_income = ""  # Set to empty string for Individual/Personal

    mobile = st.text_input("Mobile Number:")
    country = st.text_input("Country:")
    city = st.text_input("City:")
    nationality = st.text_input("Nationality:")
    cus_num = st.text_input("Customer Number (CUS NUM):")
    economic_sector = st.text_input("Economic Sector:")
    country_of_residence = st.text_input("Country of Residency:")
    resident_flag = st.text_input("Resident Flag:")
    politically_exposed_flag = st.text_input("Politically Exposed Flag (Yes/No):")
    # Use the selected customer type instead of text input
    individual_or_corporate = customer_type

    uploaded_file = st.file_uploader("Upload Excel File", type=["xls", "xlsx", "xlsm"])
    if not uploaded_file:
        st.info("Upload an Excel file to begin.")
        return

    xls = pd.ExcelFile(uploaded_file)
    first_sheet = xls.sheet_names[0]
    df = pd.read_excel(uploaded_file, sheet_name=first_sheet)

    if not set(['AMOUNT', 'TRANSACTION TYPE']).issubset(df.columns):
        st.error("The uploaded file does not contain required columns. Check your sheet format.")
        return

    debit_summary, credit_summary, all_summary = get_summaries(df)
    st.markdown("### Investigation Summaries")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("#### <span style='color:#FFE135;background:#222;border-radius:4px;padding:4px;'>Debit Transactions Summary</span>", unsafe_allow_html=True)
        st.table(style_header_and_last_row(debit_summary, "#E74C3C"))
    with col2:
        st.markdown("#### <span style='color:#44FF44;background:#222;border-radius:4px;padding:4px;'>Credit Transactions Summary</span>", unsafe_allow_html=True)
        st.table(style_header_and_last_row(credit_summary, "#27AE60"))
    with col3:
        st.markdown("#### <span style='color:#2E86C1;background:#222;border-radius:4px;padding:4px;'>All Transactions Summary</span>", unsafe_allow_html=True)
        st.table(style_header_and_last_row(all_summary, "#2E86C1"))

    st.markdown("## Top 5 Transactions for Each Transaction Type")
    if "TRANSACTION TYPE" in df.columns and "AMOUNT" in df.columns:
        df = ensure_datetime_cols(df)
        for ttype in df["TRANSACTION TYPE"].unique():
            st.markdown(f"**{ttype}**")
            ttype_df = df[df["TRANSACTION TYPE"] == ttype].copy()
            ttype_df["AMOUNT_ABS"] = ttype_df["AMOUNT"].apply(lambda x: abs(parse_amount(x)))
            top5 = ttype_df.sort_values("AMOUNT_ABS", ascending=False).head(5)
            display_cols = [col for col in ["DATE_TIME", "AMOUNT", "NAME", "BENEFICIARY", "COUNTERPARTY", "DESCRIPTION"] if col in top5.columns]
            st.dataframe(top5[display_cols] if display_cols else top5)
    else:
        st.warning("Cannot display top transactions: missing 'TRANSACTION TYPE' or 'AMOUNT' columns.")

    name_col = col_exists(df, ["NAME", "BENEFICIARY", "COUNTERPARTY"])
    st.markdown("## Top 10 Names in Transactions")
    if name_col:
        st.dataframe(top10_names_table(df))
    else:
        st.info("No name column found to display top 10 names.")

    comments = st.text_area("Analyst Comments/Findings", "")

    analyze_pressed = st.button("Analyze The Case")
    if analyze_pressed:
        if client is None:
            st.session_state["ai_summary"] = "OpenAI client not configured. Set OPENAI_API_KEY."
        else:
            with st.spinner("Contacting AI Analyst..."):
                prompt = create_ai_prompt(customer_name, job_title, organization, monthly_income,
                                          debit_summary, credit_summary, all_summary)
                try:
                    response = client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional AML compliance analyst. Respond with professional English, clear summary, and conclusion."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=800, temperature=0.2,
                    )
                    st.session_state["ai_summary"] = response.choices[0].message.content
                except Exception as e:
                    st.session_state["ai_summary"] = f"Error: {e}"

    if analyze_pressed or st.session_state.get("ai_summary"):
        st.markdown("### AI Analysis Summary")
        st.info(st.session_state.get("ai_summary", ""))

        # Minimal exports for the basic AML report (not EDD):
        def generate_docx(debit_summary, credit_summary, all_summary, comments,
                          customer_name, job_title, organization, monthly_income, ai_conclusion=None):
            doc = Document()
            doc.add_paragraph("AML Investigation Report")
            doc.add_paragraph(f"Customer: {customer_name}")
            if ai_conclusion: doc.add_paragraph(ai_conclusion)
            buf = BytesIO(); doc.save(buf); buf.seek(0); return buf

        def export_pdf_basic(debit_summary, credit_summary, all_summary, comments,
                             customer_name, job_title, organization, monthly_income):
            pdf = PDF_AR(); pdf.add_page(); pdf.write_rtl("AML Investigation Report")
            out = BytesIO(); out.write(pdf.output(dest="S").encode("latin1")); out.seek(0); return out

        cA, cB, cC = st.columns(3)
        with cA:
            if st.button("Export Report as Word (.docx)"):
                docx_file = generate_docx(debit_summary, credit_summary, all_summary, comments,
                                          customer_name, job_title, organization, monthly_income)
                st.download_button("Download Word Report", data=docx_file,
                                   file_name=f"AML_Investigation_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
        with cB:
            if st.button("Export Report as PDF"):
                pdf_file = export_pdf_basic(debit_summary, credit_summary, all_summary, comments,
                                            customer_name, job_title, organization, monthly_income)
                st.download_button("Download PDF Report", data=pdf_file, mime="application/pdf",
                                   file_name=f"AML_Investigation_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf")
        with cC:
            if st.button("Download Full Word Report with Conclusion"):
                docx_file = generate_docx(debit_summary, credit_summary, all_summary, comments,
                                          customer_name, job_title, organization, monthly_income,
                                          ai_conclusion=st.session_state.get("ai_summary", ""))
                st.download_button("Download Full Word Report", data=docx_file,
                                   file_name=f"AML_Investigation_Report_{datetime.now().strftime('%Y%m%d_%H%M')}_FULL.docx")

        st.markdown("### Next Actions")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Generate EDD for the Case (Word)"):
                edd_docx = generate_edd_docx(customer_name, job_title, organization, monthly_income,
                                             annual_salary, annual_income, mobile, country, city,
                                             nationality, cus_num, economic_sector, country_of_residence,
                                             resident_flag, politically_exposed_flag,
                                             individual_or_corporate, df=df)
                st.download_button("Download EDD Word Report", data=edd_docx,
                                   file_name=f"EDD_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
        with col2:
            if st.button("Generate EDD for the Case (PDF)"):
                edd_pdf = export_edd_pdf(customer_name, job_title, organization, monthly_income,
                                         annual_salary, annual_income, mobile, country, city,
                                         nationality, cus_num, economic_sector, country_of_residence,
                                         resident_flag, politically_exposed_flag,
                                         individual_or_corporate, df=df)
                st.download_button("Download EDD PDF Report", data=edd_pdf, mime="application/pdf",
                                   file_name=f"EDD_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf")

    # Case Control - Only for Analysts
    if current_user["role"] == "analyst":
        st.markdown("### Case Control")
        if st.button("Close the Case or Ignore the Case"):
            st.success("Case closed or ignored.")

# ----------------------------
# DASHBOARDS
# ----------------------------

def enhanced_manager_dashboard():
    st.header("🎯 Manager Command Center")
    
    # Team overview
    rows = team_leaderboard()
    st.subheader("👥 Team Overview")
    
    team_df = pd.DataFrame([{
        "User ID": r[0],
        "Username": r[1],
        "Full Name": r[2],
        "Job Title": r[3] or "",
        "Role": r[4],
        "Employee ID": r[5] or "",
        "Manager ID": r[6] or "",
        "Active": "✅ Yes" if r[7] else "❌ No",
        "Status": "🟢 Online" if is_online(r[8]) else "⚪ Offline",
        "Total Solved": int(r[9] or 0)
    } for r in rows])
    
    st.dataframe(team_df, use_container_width=True)
    
    st.divider()
    
    # Team performance chart
    st.subheader("📊 Team Performance Analytics")
    create_manager_team_chart()
    
    st.divider()
    
    # User management section
    st.subheader("⚙️ User Management")
    
    # Load user data once
    with closing(db()) as conn:
        c = conn.cursor()
        c.execute("SELECT id, username, is_active FROM users ORDER BY username")
        all_users_status = c.fetchall()
        c.execute("SELECT id, username FROM users ORDER BY username")
        all_users_simple = c.fetchall()

    tab1, tab2, tab3, tab4 = st.tabs(["➕ Create User", "🔄 Activate/Deactivate", "🔑 Reset Password", "✏️ Rename User"])

    # Create User Tab
    with tab1:
        with st.form("create_user_form"):
            new_username = st.text_input("Username").strip()
            full_name = st.text_input("Full Name").strip()
            role = st.selectbox("Role", ["analyst","manager"])
            job_title = st.text_input("Job Title (for analysts)").strip() if role=="analyst" else ""
            employee_id = st.text_input("Employee ID (for analysts)").strip() if role=="analyst" else ""
            manager_id = st.text_input("Manager ID (for managers)").strip() if role=="manager" else ""
            pwd1 = st.text_input("Password", type="password")
            pwd2 = st.text_input("Confirm Password", type="password")
            submitted = st.form_submit_button("🚀 Create User")
            
            if submitted:
                if not new_username or not full_name or not pwd1 or pwd1 != pwd2:
                    st.error("❌ Invalid input or passwords do not match.")
                else:
                    try:
                        with closing(db()) as conn:
                            c = conn.cursor()
                            c.execute("""INSERT INTO users (username, full_name, role, job_title, employee_id, manager_id, password_hash)
                                VALUES (?,?,?,?,?,?,?)""",
                                (new_username, full_name, role, job_title or None, employee_id or None, manager_id or None, bcrypt.hash(pwd1)))
                            conn.commit()
                        st.success(f"✅ User '{new_username}' created successfully!")
                    except sqlite3.IntegrityError:
                        st.error("❌ Username already exists.")

    # Activate/Deactivate Tab
    with tab2:
        target = st.selectbox("Select User",
            [f"{u[0]} - {u[1]} ({'✅ Active' if u[2] else '❌ Inactive'})" for u in all_users_status])
        if st.button("🔄 Toggle Active/Inactive"):
            user_id = int(target.split(" - ")[0])
            with closing(db()) as conn:
                c = conn.cursor()
                c.execute("UPDATE users SET is_active = 1 - is_active WHERE id=?", (user_id,))
                conn.commit()
            st.success("✅ Status updated successfully!")
            st.rerun()

    # Reset Password Tab
    with tab3:
        target2 = st.selectbox("Select User (Reset)", [f"{u[0]} - {u[1]}" for u in all_users_simple], key="reset_sel")
        newp = st.text_input("New Password", type="password")
        if st.button("🔑 Reset Password"):
            if not newp:
                st.error("❌ Password required.")
            else:
                user_id = int(target2.split(" - ")[0])
                with closing(db()) as conn:
                    c = conn.cursor()
                    c.execute("UPDATE users SET password_hash=? WHERE id=?", (bcrypt.hash(newp), user_id))
                    conn.commit()
                st.success("✅ Password reset successfully!")

    # Rename User Tab
    with tab4:
        target3 = st.selectbox("Select User to Rename", [f"{u[0]} - {u[1]}" for u in all_users_simple], key="rename_sel")
        new_username = st.text_input("New Username").strip()
        if st.button("✏️ Rename User"):
            if not new_username:
                st.error("❌ New username cannot be empty.")
            else:
                user_id = int(target3.split(" - ")[0])
                current_username = [u[1] for u in all_users_simple if u[0] == user_id][0]
                if new_username == current_username:
                    st.info("ℹ️ The new username is the same as the current one. No changes made.")
                else:
                    try:
                        with closing(db()) as conn:
                            c = conn.cursor()
                            c.execute("UPDATE users SET username=? WHERE id=?", (new_username, user_id))
                            conn.commit()
                        st.success(f"✅ Username changed from '{current_username}' to '{new_username}'.")
                        if st.session_state["auth_user"]["id"] == user_id:
                            st.session_state["auth_user"]["username"] = new_username
                            st.info("ℹ️ You renamed your own account. Session updated.")
                    except sqlite3.IntegrityError:
                        st.error("❌ That username is already taken. Choose another.")

def enhanced_analyst_dashboard():
    me = st.session_state["auth_user"]
    st.header("📊 My Analytics Dashboard")
    st.caption(f"Name: {me.get('full_name','')} | Job Title: {me.get('job_title','')} | Employee ID: {me.get('employee_id','')}")
    
    # Performance metrics
    solved, open_ = my_stats(me["id"])
    current_month, last_month, avg_monthly = get_analyst_performance_summary(me["id"])
    
    # Metrics row
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Solved", solved)
    col2.metric("Open Cases", open_)
    col3.metric("This Month", current_month, delta=current_month - last_month)
    col4.metric("Monthly Avg", avg_monthly)
    
    st.divider()
    
    # Monthly performance chart
    st.subheader("📈 Monthly Performance Trend")
    create_analyst_monthly_chart(me["id"])
    
    st.divider()
    
    # Open cases management
    st.subheader("🔄 Case Management")
    open_cases = my_open_cases(me["id"])
    if open_cases:
        options = [f"{cid} – {title or '(no title)'} (opened {ts})" for cid, title, ts in open_cases]
        sel = st.selectbox("Select a case to mark as solved", options)
        if st.button("✅ Mark Selected Case as Solved"):
            case_id = int(sel.split(" – ")[0])
            mark_case_solved(case_id)
            st.success(f"Case {case_id} marked as solved! 🎉")
            st.rerun()
    else:
        st.info("No open cases. Great work! 👍")

# ----------------------------
# START APP (Router)
# ----------------------------
init_db(); migrate_db_if_needed(); seed_manager_if_empty()

if not st.session_state["auth_user"]:
    login_block(); st.stop()

heartbeat(st.session_state["auth_user"]["id"])

with st.sidebar:
    st.write(f"👤 {st.session_state['auth_user']['username']} ({st.session_state['auth_user']['role']})")
    if st.button("Logout"):
        st.session_state["auth_user"] = None; st.rerun()

role = st.session_state["auth_user"]["role"]
page = st.sidebar.radio("Navigation", ["Dashboard", "Operations"])

if page == "Dashboard":
    enhanced_manager_dashboard() if role == "manager" else enhanced_analyst_dashboard()
elif page == "Operations":
    operations_page()

st.markdown("---\n*No data is stored. Close the tab to clear everything. — Credit: Blackyotta Company*")
